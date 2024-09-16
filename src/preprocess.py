import os
import sys
# Directly set the project root directory
# sys.path.append("E:/Training/Atomcamp/DS6_Bootcamp/Projects/FYP/fyp-chatbot")
project_root = "E:/Training/Atomcamp/DS6_Bootcamp/Projects/FYP/fyp-chatbot"
# Ensure the project root is at the top of sys.path
sys.path.insert(0, project_root)
# from langchain.document_loaders import PyPDFLoader
# from langchain_community.document_loaders import PyPDFLoader
from langchain.schema import Document
from docx import Document as DocxDocument  # Import python-docx Document for reading .docx files
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
from custom_logging import logger
from custom_exception import CustomException

# Function to extract text from PDFs
def extract_text_from_pdf(pdf_path):
    pdf_text = ""
    with open(pdf_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page_num in range(len(reader.pages)):
            pdf_text += reader.pages[page_num].extract_text()
    return pdf_text

# Function to load and extract text from different document types
def load_documents(directory_path: str):
    """
    Load documents from a PDF file.

    Args:
        directory_path (str): Path to the directory.

    Returns:
        list: List of documents.
    """
    try:
        logger.info("Loading Documents ...")
        # loader = PyPDFLoader(file_path)
        # documents = loader.load()
        
        documents = []

        # Iterate over files in the directory
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)

            # Handling text files (.txt)
            if filename.endswith(".txt"):
                with open(file_path, "r") as file:
                    content = file.read()
                    doc = Document(page_content=content, metadata={"filename": filename})
                    documents.append(doc)

            # Handling Word documents (.docx)
            elif filename.endswith(".docx"):
                # Open and read the .docx file
                docx_doc = DocxDocument(file_path)
                content = []
                
                # Extract text from the Word document
                for para in docx_doc.paragraphs:
                    content.append(para.text)
                
                # Join the content list into a single string
                content = "\n".join(content)
                
                # Create a LangChain Document object with the content and filename metadata
                doc = Document(page_content=content, metadata={"filename": filename})
                documents.append(doc)

            # Handling PDF files (.pdf)
            elif filename.endswith(".pdf"):
                pdf_text = extract_text_from_pdf(file_path)
                doc = Document(page_content=pdf_text, metadata={"filename": filename})
                documents.append(doc)

        logger.info("Documents loaded successfully from %s", file_path)

        return documents
    
    except Exception as e:
        raise CustomException(e, sys)

# Function to split the documents into chunks
def split_documents(documents: list, chunk_size: int = 1500, chunk_overlap: int = 250):
    """
    Split documents into smaller chunks.

    Args:
        documents (list): List of documents.
        chunk_size (int): Size of each chunk.
        chunk_overlap (int): Overlap between chunks.

    Returns:
        list: List of text chunks.
    """
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        texts = text_splitter.split_documents(documents)
        logger.info("Documents split into chunks successfully")

        return texts
    
    except Exception as e:
        raise CustomException(e, sys)
    
    
if __name__ == "__main__":
    directory = "E:/Training/Atomcamp/DS6_Bootcamp/Projects/FYP/Rules_and_Policies"
    directory_path = os.path.join(project_root, directory)
    documents = load_documents(directory_path)
    texts = split_documents(documents)
    
    # print('*' * 60)
    # print(len(texts))
    # print(len(documents))
    # print('*' * 60)